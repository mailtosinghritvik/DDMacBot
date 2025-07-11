import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import time
import os
import tempfile
from openai import OpenAI

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Set page configuration
st.set_page_config(
    page_title="Document Generation - DDMac Bot",
    page_icon="📄",
    layout="wide"
)

# Document Generation Functions
def generate_proposal_document():
    """Generate a proposal document using the assistant with code interpreter and template"""
    try:
        if not st.session_state.assistant_id or not st.session_state.thread_id:
            return None, "No assistant found. Please upload an Excel file first on the Home page."
        
        # Get project and sheet info from session state
        project_info = st.session_state.get('project_info', {})
        sheet_info = st.session_state.get('sheet_info', {})
        
        # Get user customizations
        customizations = st.session_state.get('proposal_customizations', {
            'excel_interpretations': {},
            'cost_mappings': {},
            'template_replacements': {},
            'custom_instructions': "",
            'exclude_sheets': [],
            'cost_categories': {}
        })
        
        # Create comprehensive project context
        project_context = f"""
PROJECT INFORMATION:
- Company Name: {project_info.get('company_name', 'N/A')}
- Project Name: {project_info.get('project_name', 'N/A')}
- Project Type: {project_info.get('project_type', 'N/A')}
- Location: {project_info.get('project_location', 'N/A')}
- Contract Terms: {project_info.get('contract_terms', 'N/A')}
- Additional Info: {project_info.get('additional_info', 'N/A')}

EXCEL SHEET INFORMATION:
"""
        
        # Add detailed sheet information with custom interpretations
        for sheet_name, info in sheet_info.items():
            # Skip excluded sheets
            if sheet_name in customizations.get('exclude_sheets', []):
                continue
                
            project_context += f"""
Sheet: {sheet_name}
- Purpose/Meaning: {info.get('meaning', 'N/A')}
- Description: {info.get('description', 'N/A')}"""
            
            # Add custom interpretations if available
            custom_desc = customizations['excel_interpretations'].get(f"{sheet_name}_description", "")
            if custom_desc:
                project_context += f"\n- Custom Description: {custom_desc}"
            
            custom_cols = customizations['excel_interpretations'].get(f"{sheet_name}_columns", "")
            if custom_cols:
                project_context += f"\n- Column Information: {custom_cols}"
            
            cost_handling = customizations['cost_mappings'].get(sheet_name, "")
            if cost_handling:
                project_context += f"\n- Cost Handling: {cost_handling}"
            
            project_context += "\n"
        
        # Add custom instructions if provided
        if customizations.get('custom_instructions'):
            project_context += f"""

CUSTOM INSTRUCTIONS:
{customizations['custom_instructions']}
"""
        
        # Upload proposal template DOCX file for this session
        template_file_obj = None
        template_path = "/Users/ritviksingh/Desktop/Ace148/DDMacBot/resources/prompts/ddmac_template.docx"
        if os.path.exists(template_path):
            try:
                with open(template_path, "rb") as f:
                    template_file_obj = client.files.create(
                        file=f,
                        purpose='assistants'
                    )
                
                # Update assistant with the template file
                current_assistant = client.beta.assistants.retrieve(st.session_state.assistant_id)
                current_file_ids = current_assistant.tool_resources.code_interpreter.file_ids
                updated_file_ids = list(current_file_ids) + [template_file_obj.id]
                
                client.beta.assistants.update(
                    assistant_id=st.session_state.assistant_id,
                    tool_resources={
                        "code_interpreter": {
                            "file_ids": updated_file_ids
                        }
                    }
                )
                
            except Exception as template_error:
                st.warning(f"Could not upload proposal template: {str(template_error)}")
        
        # Create the special prompt based on your specifications
        prompt = f"""{project_context}

TASK: Generate a professional electrical project proposal document using the provided ddmac_template.docx file.

You have access to:
1. An Excel file with project cost data and estimates
2. A ddmac_template.docx file with specific tags to replace

CRITICAL REQUIREMENTS:
- You MUST generate and save a DOCX file for download
- Use the python-docx library to modify the template
- Save the file with a clear filename ending in .docx
- Provide the saved file for download

STEP-BY-STEP INSTRUCTIONS:

STEP 1: ACCESS FILES CORRECTLY
First, list all available files and identify them:

```python
import os
import pandas as pd
from docx import Document

print("=== AVAILABLE FILES ===")
files = os.listdir('.')
for file in files:
    size = os.path.getsize(file)
    print(f"File: {{file}} - {{size}} bytes")

# The Excel file should be one of these files (look for the smaller one, typically 5-50KB)
# The DOCX template should be the larger file (around 1MB)
```

STEP 2: READ THE EXCEL FILE
Try to read the Excel file by its actual filename in the environment:

```python
# Find and read the Excel file
excel_file = None
for file in os.listdir('.'):
    try:
        # Try to read as Excel - if successful, this is our Excel file
        df = pd.read_excel(file)
        excel_file = file
        print(f"✅ Found Excel file: {{file}}")
        print(f"Sheets available: {{pd.ExcelFile(file).sheet_names}}")
        break
    except:
        continue

if excel_file:
    # Read all sheets
    excel_data = pd.ExcelFile(excel_file)
    all_sheets = {{}}
    for sheet_name in excel_data.sheet_names:
        sheet_df = pd.read_excel(excel_data, sheet_name=sheet_name)
        all_sheets[sheet_name] = sheet_df
        print(f"Sheet '{{sheet_name}}': {{sheet_df.shape[0]}} rows, {{sheet_df.shape[1]}} columns")
        print(sheet_df.head())
```

STEP 3: READ THE DOCX TEMPLATE
Find and read the DOCX template:

```python
# Find and read the DOCX template
template_file = None
for file in os.listdir('.'):
    try:
        # Try to read as DOCX - if successful, this is our template
        doc = Document(file)
        template_file = file
        print(f"✅ Found DOCX template: {{file}}")
        print(f"Number of paragraphs: {{len(doc.paragraphs)}}")
        break
    except:
        continue
```

STEP 4: REPLACE TEMPLATE TAGS
Replace the following tags with data from Excel and project info:

INSERT_TO_COMPANY, INSERT_NAME_OF_PERSON, INSERT_PROJECT_NAME, DATE, EST, INSERT_ON_THE_BASIS_PARAGRAPH_HERE, INSERT_COST_ONE_NAME, INSERT_COST_ONE_VALUE, INSERT_COST_TWO_NAME, INSERT_COST_TWO_VALUE, INSERT_COST_THREE_NAME, INSERT_COST_THREE_VALUE, INSERT_COST_FOUR_NAME, INSERT_COST_FOUR_VALUE, PROJECT_ORGANISER_NAME, PROJECT_ORGANISER_TITLE, ITEM_NOT_INCLUDED_1, ITEM_NOT_INCLUDED_2, QUALIFICATION_1, QUALIFICATION_2

SPECIFIC MAPPINGS:
- INSERT_TO_COMPANY: {customizations['template_replacements'].get('INSERT_TO_COMPANY', project_info.get('company_name', 'Client Company'))}
- INSERT_NAME_OF_PERSON: {customizations['template_replacements'].get('INSERT_NAME_OF_PERSON', 'Project Manager or Contact Person')}
- INSERT_PROJECT_NAME: {customizations['template_replacements'].get('INSERT_PROJECT_NAME', project_info.get('project_name', 'Electrical Project'))}
- DATE: {datetime.now().strftime('%B %d, %Y')}
- EST: Generate appropriate estimate number
- INSERT_ON_THE_BASIS_PARAGRAPH_HERE: {customizations['template_replacements'].get('INSERT_ON_THE_BASIS_PARAGRAPH_HERE', 'Create description based on project type and Excel analysis')}
- PROJECT_ORGANISER_NAME: {customizations['template_replacements'].get('PROJECT_ORGANISER_NAME', project_info.get('company_name', 'DDMac Electrical'))}
- PROJECT_ORGANISER_TITLE: {customizations['template_replacements'].get('PROJECT_ORGANISER_TITLE', 'Electrical Contractor')}

COST CATEGORIES TO USE:
- Cost One: {customizations['cost_categories'].get('COST_ONE', 'Labor Costs')}
- Cost Two: {customizations['cost_categories'].get('COST_TWO', 'Material Costs')}
- Cost Three: {customizations['cost_categories'].get('COST_THREE', 'Equipment Costs')}
- Cost Four: {customizations['cost_categories'].get('COST_FOUR', 'Other Costs')}

```python
# Extract important cost data from Excel sheets
major_costs = []
for sheet_name, df in all_sheets.items():
    # Look for cost-related columns
    cost_columns = [col for col in df.columns if any(keyword in col.lower() for keyword in ['cost', 'total', 'price', 'amount'])]
    
    for cost_col in cost_columns:
        if df[cost_col].dtype in ['float64', 'int64']:
            total_cost = df[cost_col].sum()
            major_costs.append((f"{{sheet_name}} - {{cost_col}}", total_cost))

# Sort by cost amount to get the most significant ones
major_costs.sort(key=lambda x: x[1], reverse=True)

# Use top 4 costs for the template with custom category names
cost_one_name = major_costs[0][0] if len(major_costs) > 0 else "{customizations['cost_categories'].get('COST_ONE', 'Labor Costs')}"
cost_one_value = f"${{major_costs[0][1]:,.2f}}" if len(major_costs) > 0 else "$0.00"
cost_two_name = major_costs[1][0] if len(major_costs) > 1 else "{customizations['cost_categories'].get('COST_TWO', 'Material Costs')}"  
cost_two_value = f"${{major_costs[1][1]:,.2f}}" if len(major_costs) > 1 else "$0.00"
cost_three_name = major_costs[2][0] if len(major_costs) > 2 else "{customizations['cost_categories'].get('COST_THREE', 'Equipment Costs')}"
cost_three_value = f"${{major_costs[2][1]:,.2f}}" if len(major_costs) > 2 else "$0.00"
cost_four_name = major_costs[3][0] if len(major_costs) > 3 else "{customizations['cost_categories'].get('COST_FOUR', 'Other Costs')}"
cost_four_value = f"${{major_costs[3][1]:,.2f}}" if len(major_costs) > 3 else "$0.00"

# Apply any custom cost calculation instructions
{f"# Custom cost calculation instructions: {customizations.get('cost_calculation_instructions', '')}" if customizations.get('cost_calculation_instructions') else "# No custom cost calculation instructions"}

# Replace tags in document
replacements = {{
    'INSERT_TO_COMPANY': '{customizations['template_replacements'].get('INSERT_TO_COMPANY', project_info.get('company_name', 'Client Company'))}',
    'INSERT_NAME_OF_PERSON': '{customizations['template_replacements'].get('INSERT_NAME_OF_PERSON', 'Project Manager')}',
    'INSERT_PROJECT_NAME': '{customizations['template_replacements'].get('INSERT_PROJECT_NAME', project_info.get('project_name', 'Electrical Project'))}',
    'DATE': '{datetime.now().strftime('%B %d, %Y')}',
    'EST': 'EST-2025-001',
    'INSERT_ON_THE_BASIS_PARAGRAPH_HERE': '{customizations['template_replacements'].get('INSERT_ON_THE_BASIS_PARAGRAPH_HERE', 'Based on our analysis of the provided project specifications and cost data.')}',
    'INSERT_COST_ONE_NAME': cost_one_name,
    'INSERT_COST_ONE_VALUE': cost_one_value,
    'INSERT_COST_TWO_NAME': cost_two_name,
    'INSERT_COST_TWO_VALUE': cost_two_value,
    'INSERT_COST_THREE_NAME': cost_three_name,
    'INSERT_COST_THREE_VALUE': cost_three_value,
    'INSERT_COST_FOUR_NAME': cost_four_name,
    'INSERT_COST_FOUR_VALUE': cost_four_value,
    'PROJECT_ORGANISER_NAME': '{customizations['template_replacements'].get('PROJECT_ORGANISER_NAME', project_info.get('company_name', 'DDMac Electrical'))}',
    'PROJECT_ORGANISER_TITLE': '{customizations['template_replacements'].get('PROJECT_ORGANISER_TITLE', 'Electrical Contractor')}',
    'ITEM_NOT_INCLUDED_1': '{customizations['template_replacements'].get('ITEM_NOT_INCLUDED_1', 'Permits and inspections')}',
    'ITEM_NOT_INCLUDED_2': '{customizations['template_replacements'].get('ITEM_NOT_INCLUDED_2', 'Site preparation')}',
    'QUALIFICATION_1': '{customizations['template_replacements'].get('QUALIFICATION_1', 'All work to be performed according to local electrical codes')}',
    'QUALIFICATION_2': '{customizations['template_replacements'].get('QUALIFICATION_2', 'Material costs subject to change based on availability')}'
}}

# Apply replacements to document
for paragraph in doc.paragraphs:
    for tag, replacement in replacements.items():
        if tag in paragraph.text:
            paragraph.text = paragraph.text.replace(tag, replacement)

# Also check tables for tags
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for tag, replacement in replacements.items():
                if tag in cell.text:
                    cell.text = cell.text.replace(tag, replacement)
```

STEP 5: SAVE THE FINAL DOCUMENT
```python
# Save the modified document
output_file = "DDMac_Proposal_Final.docx"
doc.save(output_file)
print(f"✅ Proposal document saved successfully as: {{output_file}}")

# Verify file was created
if os.path.exists(output_file):
    size = os.path.getsize(output_file)
    print(f"✅ File verified: {{output_file}} ({{size}} bytes)")
else:
    print("❌ Error: File was not created")
```

CRITICAL: You MUST save the file with the exact filename "DDMac_Proposal_Final.docx" for download."""

        # Send the message to the assistant
        message = client.beta.threads.messages.create(
            thread_id=st.session_state.thread_id,
            role="user",
            content=prompt
        )
        
        # Create and run the assistant
        run = client.beta.threads.runs.create(
            thread_id=st.session_state.thread_id,
            assistant_id=st.session_state.assistant_id
        )
        
        return run.id, "Proposal generation started successfully"
        
    except Exception as e:
        return None, f"Error generating proposal: {str(e)}"

def check_proposal_generation_status(run_id):
    """Check the status of proposal generation and retrieve the document if ready"""
    try:
        run = client.beta.threads.runs.retrieve(
            thread_id=st.session_state.thread_id,
            run_id=run_id
        )
        
        if run.status == "completed":
            # Retrieve the latest message from the assistant
            messages = client.beta.threads.messages.list(
                thread_id=st.session_state.thread_id,
                limit=1
            )
            
            if messages.data:
                message = messages.data[0]
                
                # Check for file annotations in the message content
                files_found = []
                for content_block in message.content:
                    if hasattr(content_block, 'text') and hasattr(content_block.text, 'annotations'):
                        for annotation in content_block.text.annotations:
                            if hasattr(annotation, 'file_path'):
                                files_found.append(annotation.file_path.file_id)
                            elif hasattr(annotation, 'file_download'):
                                files_found.append(annotation.file_download.file_id)
                
                # Also check for any file attachments in the run steps
                if not files_found:
                    try:
                        run_steps = client.beta.threads.runs.steps.list(
                            thread_id=st.session_state.thread_id,
                            run_id=run_id
                        )
                        
                        for step in run_steps.data:
                            if hasattr(step, 'step_details') and hasattr(step.step_details, 'tool_calls'):
                                for tool_call in step.step_details.tool_calls:
                                    if hasattr(tool_call, 'code_interpreter') and hasattr(tool_call.code_interpreter, 'outputs'):
                                        for output in tool_call.code_interpreter.outputs:
                                            if hasattr(output, 'logs'):
                                                # Check logs for file creation messages
                                                if "Document saved as:" in output.logs or ".docx" in output.logs:
                                                    # Try to find file IDs in the logs
                                                    continue
                    except Exception as step_error:
                        print(f"Error checking run steps: {step_error}")
                
                if files_found:
                    return "completed", files_found[0], "Proposal document generated successfully"
                else:
                    # Check message content for any useful info
                    content = message.content[0].text.value if message.content else "No response"
                    return "completed", None, f"Generation completed but no file found. Response: {content[:300]}..."
            else:
                return "completed", None, "No response from assistant"
                
        elif run.status == "failed":
            error_msg = run.last_error.message if run.last_error else 'Unknown error'
            return "failed", None, f"Proposal generation failed: {error_msg}"
        elif run.status in ["queued", "in_progress", "requires_action"]:
            return "in_progress", None, "Proposal generation in progress..."
        else:
            return "unknown", None, f"Unknown status: {run.status}"
            
    except Exception as e:
        return "error", None, f"Error checking status: {str(e)}"

def download_generated_file(file_id, default_filename="DDMac_Proposal.docx"):
    """Download the generated file from OpenAI"""
    try:
        # Get file content from OpenAI
        file_content = client.files.content(file_id)
        file_bytes = file_content.read()
        
        # Try to get original filename, fallback to default
        try:
            file_info = client.files.retrieve(file_id)
            filename = file_info.filename or default_filename
        except:
            filename = default_filename
            
        return file_bytes, filename
        
    except Exception as e:
        st.error(f"Error downloading file: {str(e)}")
        return None, None

# Custom CSS
st.markdown("""
<style>
    .main-header {
        text-align: center;
        padding: 2rem 0;
        background: linear-gradient(90deg, #1f4e79 0%, #2e7d32 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    
    .document-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        margin: 1rem 0;
        border-left: 4px solid #1f4e79;
    }
    
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 10px;
        text-align: center;
        margin: 0.5rem 0;
    }
    
    .generation-status {
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        border-left: 4px solid #28a745;
    }
    
    .status-pending { background-color: #fff3cd; color: #856404; border-left-color: #ffc107; }
    .status-processing { background-color: #cce5ff; color: #004085; border-left-color: #007bff; }
    .status-complete { background-color: #d4edda; color: #155724; border-left-color: #28a745; }
    .status-error { background-color: #f8d7da; color: #721c24; border-left-color: #dc3545; }
    
    .project-info {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }
    
    .template-info {
        background: #e3f2fd;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        border: 1px solid #bbdefb;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state for document generation
if 'generating_proposal' not in st.session_state:
    st.session_state.generating_proposal = False

# Header
st.markdown("""
<div class="main-header">
    <h1>📄 Document Generation Center</h1>
    <p>Generate professional electrical project proposals and documents using AI</p>
</div>
""", unsafe_allow_html=True)

# Check if we have the necessary session state data
assistant_id = st.session_state.get('assistant_id')
thread_id = st.session_state.get('thread_id')

if not assistant_id or not thread_id:
    st.warning("⚠️ No active assistant found. Please upload and process an Excel file on the Home page first.")
    st.info("👉 Go to the Home page to upload your AccuBid Excel file and create a dedicated assistant.")
    
    # Debug information
    with st.expander("🔍 Debug Information"):
        st.write("**Current Session State:**")
        st.write(f"- Assistant ID: {assistant_id or 'Not set'}")
        st.write(f"- Thread ID: {thread_id or 'Not set'}")
        st.write(f"- Project Info: {'Available' if st.session_state.get('project_info') else 'Not available'}")
        st.write(f"- Sheet Info: {'Available' if st.session_state.get('sheet_info') else 'Not available'}")
        st.write(f"- Generating Proposal: {st.session_state.get('generating_proposal', 'Not set')}")
        
        if st.button("🔄 Clear All Session State (for testing)"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.success("Session state cleared!")
            st.rerun()
    
    st.stop()

# Project information display
if st.session_state.get('project_info'):
    project_info = st.session_state.project_info
    st.markdown("""
    <div class="document-card">
        <h3>📋 Current Project Information</h3>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Company", project_info.get('company_name', 'N/A'))
        st.metric("Project Type", project_info.get('project_type', 'N/A'))
    with col2:
        st.metric("Project Name", project_info.get('project_name', 'N/A'))
        st.metric("Location", project_info.get('project_location', 'N/A') or 'Not specified')
    with col3:
        st.metric("Assistant ID", f"...{st.session_state.assistant_id[-8:]}")
        st.metric("Thread ID", f"...{st.session_state.thread_id[-8:]}")

# Document Generation Section
st.markdown("""
<div class="document-card">
    <h3>📝 Professional Proposal Generation</h3>
    <p>Generate a comprehensive electrical project proposal using your Excel data and our professional template.</p>
</div>
""", unsafe_allow_html=True)

# Template information
st.markdown("""
<div class="template-info">
    <h4>📄 Template Information</h4>
    <p><strong>Template:</strong> DDMac Professional Proposal Template</p>
    <p><strong>Features:</strong> Automated cost integration, professional formatting, customizable sections</p>
    <p><strong>Output:</strong> Microsoft Word document (.docx) ready for client presentation</p>
</div>
""", unsafe_allow_html=True)

# Document generation controls
if not st.session_state.get('generating_proposal', False):
    # Check if user has customizations
    has_customizations = any([
        st.session_state.get('proposal_customizations', {}).get('excel_interpretations'),
        st.session_state.get('proposal_customizations', {}).get('template_replacements'),
        st.session_state.get('proposal_customizations', {}).get('custom_instructions'),
        st.session_state.get('proposal_customizations', {}).get('cost_categories')
    ])
    
    if not has_customizations:
        st.markdown("""
        <div class="template-info">
            <h4>💡 Enhance Your Proposal</h4>
            <p><strong>Want better results?</strong> Customize how the AI interprets your Excel data and generates your proposal!</p>
            <p>• Explain what your Excel columns mean<br>
            • Define custom cost categories<br>
            • Set template preferences<br>
            • Add special instructions</p>
        </div>
        """, unsafe_allow_html=True)
        
        col_cust1, col_cust2 = st.columns([1, 1])
        with col_cust1:
            if st.button("⚙️ Customize Proposal Settings", use_container_width=True, key="go_to_customize"):
                st.switch_page("pages/Customize_Proposal.py")
        with col_cust2:
            st.info("Recommended for best results!")
    else:
        st.success("✅ Using your custom proposal settings!")
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        generate_button = st.button("🚀 Generate Proposal Document", type="primary", use_container_width=True)
        if generate_button:
            # Debug: Show that button was clicked
            st.info("Button clicked! Initializing proposal generation...")
            st.session_state.generating_proposal = True
            # Small delay to ensure user sees the feedback
            time.sleep(0.5)
            st.rerun()
    
    with col2:
        st.markdown("""
        **What this will do:**
        - Analyze your Excel data for cost information
        - Replace template placeholders with project-specific data
        - Generate a professional Word document
        - Provide download link for the completed proposal
        """)

# Handle Proposal Generation
if st.session_state.get('generating_proposal', False):
    st.markdown("---")
    st.success("🎯 Proposal generation mode activated!")
    
    # Initialize the generation if not already started
    if 'proposal_run_id' not in st.session_state:
        with st.spinner("🚀 Starting proposal generation..."):
            run_id, message = generate_proposal_document()
            if run_id:
                st.session_state.proposal_run_id = run_id
                st.success(f"✅ {message}")
                st.info(f"📋 Run ID: {run_id}")
                st.markdown("""
                <div class="generation-status status-processing">
                    <h4>⏳ Generation in Progress</h4>
                    <p>The assistant is analyzing your Excel data and generating the proposal document...</p>
                </div>
                """, unsafe_allow_html=True)
                st.rerun()
            else:
                st.error(f"❌ {message}")
                st.session_state.generating_proposal = False
                st.rerun()
    else:
        # Check the status of the ongoing generation
        status, file_id, message = check_proposal_generation_status(st.session_state.proposal_run_id)
        
        if status == "completed":
            if file_id:
                st.markdown("""
                <div class="generation-status status-complete">
                    <h4>🎉 Proposal Generated Successfully!</h4>
                    <p>Your professional proposal has been created with data from your Excel file!</p>
                </div>
                """, unsafe_allow_html=True)
                
                # Download the file
                file_content, filename = download_generated_file(file_id, "DDMac_Proposal.docx")
                if file_content:
                    col1, col2 = st.columns([1, 1])
                    with col1:
                        st.download_button(
                            label="📥 Download Proposal Document",
                            data=file_content,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            type="primary"
                        )
                    with col2:
                        st.info("💡 The proposal includes cost analysis from your Excel data and follows professional formatting standards.")
                
                # Clean up session state
                st.markdown("---")
                col_done1, col_done2, col_done3 = st.columns(3)
                with col_done1:
                    if st.button("✅ Done", use_container_width=True, key="done_button_success"):
                        st.session_state.generating_proposal = False
                        if 'proposal_run_id' in st.session_state:
                            del st.session_state.proposal_run_id
                        st.rerun()
                with col_done2:
                    if st.button("🔄 Generate Another", use_container_width=True, key="generate_another_button"):
                        st.session_state.generating_proposal = False
                        if 'proposal_run_id' in st.session_state:
                            del st.session_state.proposal_run_id
                        st.session_state.generating_proposal = True
                        st.rerun()
                with col_done3:
                    if st.button("🏠 Back to Home", use_container_width=True, key="back_home_success"):
                        st.session_state.generating_proposal = False
                        if 'proposal_run_id' in st.session_state:
                            del st.session_state.proposal_run_id
                        st.switch_page("Home.py")
            else:
                st.markdown("""
                <div class="generation-status status-error">
                    <h4>⚠️ Generation Completed but No Document Produced</h4>
                    <p>The assistant finished processing but didn't create a downloadable file.</p>
                </div>
                """, unsafe_allow_html=True)
                
                with st.expander("🔍 View Assistant Response"):
                    st.text(message)
                
                if st.button("🔄 Try Again", use_container_width=True, key="try_again_1"):
                    st.session_state.generating_proposal = False
                    if 'proposal_run_id' in st.session_state:
                        del st.session_state.proposal_run_id
                    st.rerun()
                    
        elif status == "failed":
            st.markdown("""
            <div class="generation-status status-error">
                <h4>❌ Generation Failed</h4>
                <p>The assistant encountered an issue while generating the proposal.</p>
            </div>
            """, unsafe_allow_html=True)
            
            st.error(f"Error: {message}")
            
            if st.button("🔄 Try Again", use_container_width=True, key="try_again_2"):
                st.session_state.generating_proposal = False
                if 'proposal_run_id' in st.session_state:
                    del st.session_state.proposal_run_id
                st.rerun()
                
        elif status == "in_progress":
            st.markdown("""
            <div class="generation-status status-processing">
                <h4>⏳ Processing Your Request</h4>
                <p>The assistant is working on your proposal... This may take a few moments as it analyzes your Excel data and formats the document.</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Progress indicator
            progress_bar = st.progress(0)
            for i in range(100):
                progress_bar.progress(i + 1)
                time.sleep(0.02)
            
            # Auto-refresh every 3 seconds
            time.sleep(1)
            st.rerun()
            
        else:
            st.markdown("""
            <div class="generation-status status-error">
                <h4>⚠️ Unknown Status</h4>
                <p>Encountered an unexpected status during generation.</p>
            </div>
            """, unsafe_allow_html=True)
            
            st.error(f"Status: {message}")
            
            if st.button("🔄 Try Again", use_container_width=True, key="try_again_3"):
                st.session_state.generating_proposal = False
                if 'proposal_run_id' in st.session_state:
                    del st.session_state.proposal_run_id
                st.rerun()

# Future Features Section
st.markdown("---")
st.markdown("""
<div class="document-card">
    <h3>🚀 Coming Soon</h3>
    <p>Additional document generation features in development:</p>
</div>
""", unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("""
    <div class="template-info">
        <h4>📊 Custom Reports</h4>
        <p>Generate detailed cost analysis reports, material breakdowns, and labor summaries</p>
        <small>Status: In Development</small>
    </div>
    """, unsafe_allow_html=True)

with col2:
    st.markdown("""
    <div class="template-info">
        <h4>📧 Email Templates</h4>
        <p>Professional email templates for client communication and project updates</p>
        <small>Status: Planned</small>
    </div>
    """, unsafe_allow_html=True)

with col3:
    st.markdown("""
    <div class="template-info">
        <h4>📋 Multiple Templates</h4>
        <p>Choose from various proposal templates for different project types</p>
        <small>Status: Planned</small>
    </div>
    """, unsafe_allow_html=True)

# Help Section
with st.sidebar:
    st.header("📋 Generation Guide")
    
    st.markdown("""
    ### How It Works
    1. **Upload Excel File** on Home page
    2. **Process with AI** to create assistant
    3. **Customize Settings** (recommended)
    4. **Generate Documents** here
    5. **Download and Use** with clients
    
    ### Template Features
    - **Automated Costs**: Pulls data from Excel
    - **Professional Format**: Industry-standard layout
    - **Customizable**: Adapts to project type
    - **Client-Ready**: Professional presentation
    
    ### Troubleshooting
    - Ensure Excel file is processed first
    - Check that assistant is active
    - Verify template file exists
    - Try regenerating if errors occur
    """)
    
    if st.button("⚙️ Customize Proposal", use_container_width=True, key="sidebar_customize"):
        st.switch_page("pages/Customize_Proposal.py")
    
    if st.button("🏠 Back to Home", use_container_width=True, key="back_home_sidebar"):
        st.switch_page("Home.py")
